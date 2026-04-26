"""
ClinicalTable Pro — DOCX Export (python-docx)

Generates a Word document with a journal-formatted Table 1.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def export_docx(table_data: dict[str, Any], title: str = "Table 1") -> bytes:
    """
    Generate DOCX bytes from table data.

    Parameters
    ----------
    table_data : dict
        Output from generate_table_one().
    title : str
        Table title/number.

    Returns
    -------
    bytes : DOCX file content.
    """

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.orientation = 1  # Landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{title}. Baseline Characteristics of Study Participants")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = "Times New Roman"
    title_para.space_after = Pt(6)

    columns = table_data["columns"]
    rows = table_data["rows"]
    footnotes = table_data.get("footnotes", [])
    stat_notes = table_data.get("stats_method_notes", [])

    # Create table
    n_cols = len(columns)
    n_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style: no default borders
    table.style = "Table Grid"

    # Remove all borders first, then add journal-style ones
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={"val": "none"},
                bottom={"val": "none"},
                left={"val": "none"},
                right={"val": "none"},
            )

    # Header row
    header_row = table.rows[0]
    for i, col in enumerate(columns):
        cell = header_row.cells[i]
        cell.text = col
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

        # Top and bottom border for header
        _set_cell_border(
            cell,
            top={"val": "single", "sz": "12", "color": "000000"},
            bottom={"val": "single", "sz": "6", "color": "000000"},
        )

    # Data rows
    for r_idx, row_data in enumerate(rows):
        doc_row = table.rows[r_idx + 1]
        row_type = row_data.get("type", "")
        variable = row_data.get("variable", "")
        values = row_data.get("values", {})

        # Variable name cell
        cell = doc_row.cells[0]
        if row_type == "category_value":
            cell.text = f"    {variable.strip()}"
        else:
            cell.text = variable
        para = cell.paragraphs[0]
        if para.runs:
            run = para.runs[0]
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            if row_type == "category_header":
                run.bold = True
                run.italic = True

        # Value cells
        for c_idx, col in enumerate(columns[1:], 1):
            cell = doc_row.cells[c_idx]
            if col == "P-value":
                val = values.get("p_value", "")
            else:
                group_name = col.split(" (n=")[0] if " (n=" in col else col.split(" (N=")[0]
                val = values.get(group_name, values.get(col, ""))
            cell.text = str(val)
            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

        # Bottom border on last row
        if r_idx == len(rows) - 1:
            for cell in doc_row.cells:
                _set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "12", "color": "000000"},
                )

    # Footnotes
    all_notes = footnotes + stat_notes
    if all_notes:
        footnote_para = doc.add_paragraph()
        footnote_para.space_before = Pt(4)
        for note in all_notes:
            run = footnote_para.add_run(note + "\n")
            run.font.size = Pt(8)
            run.font.name = "Times New Roman"
            run.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
